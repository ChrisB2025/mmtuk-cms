"""
DOCX extraction service: text and image extraction using python-docx.
"""

import io
import logging

from docx import Document
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from PIL import Image

logger = logging.getLogger(__name__)

MAX_FILE_SIZE = 20 * 1024 * 1024  # 20 MB
MAX_IMAGES = 20
MIN_IMAGE_SIZE = 50  # px


def extract_docx(file_bytes, filename):
    """
    Extract text and images from a Word (.docx) file.

    Args:
        file_bytes: Raw bytes of the .docx file.
        filename: Original filename (for logging/metadata).

    Returns:
        dict with keys: filename, page_count, text, images, metadata.
        Each image is {page, data, ext, width, height}.

    Raises:
        ValueError: If the file is too large or cannot be opened.
    """
    if len(file_bytes) > MAX_FILE_SIZE:
        raise ValueError(f'File exceeds the {MAX_FILE_SIZE // (1024 * 1024)}MB size limit.')

    try:
        doc = Document(io.BytesIO(file_bytes))
    except Exception as exc:
        raise ValueError(f'Could not open document: {exc}')

    # Extract metadata
    props = doc.core_properties
    metadata = {
        'title': props.title or '',
        'author': props.author or '',
        'subject': props.subject or '',
    }

    # Extract text from paragraphs
    text_parts = [p.text for p in doc.paragraphs if p.text.strip()]
    # Also extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
            if row_text:
                text_parts.append(row_text)

    full_text = '\n\n'.join(text_parts)

    # Approximate page count (roughly 40 paragraphs per page)
    page_count = max(1, len(doc.paragraphs) // 40)

    # Extract images from document relationships
    images = []
    try:
        for rel in doc.part.rels.values():
            if len(images) >= MAX_IMAGES:
                break
            if rel.reltype == RT.IMAGE:
                try:
                    img_bytes = rel.target_part.blob
                    img_ext = rel.target_part.content_type.split('/')[-1].lower()
                    # Normalise extension
                    if img_ext in ('jpeg', 'jpg'):
                        img_ext = 'jpeg'
                    elif img_ext == 'png':
                        img_ext = 'png'
                    else:
                        img_ext = 'png'

                    # Get dimensions and skip tiny images
                    try:
                        with Image.open(io.BytesIO(img_bytes)) as pil_img:
                            w, h = pil_img.size
                    except Exception:
                        continue

                    if w < MIN_IMAGE_SIZE or h < MIN_IMAGE_SIZE:
                        continue

                    images.append({
                        'page': 1,  # docx doesn't have true page numbers
                        'data': img_bytes,
                        'ext': img_ext,
                        'width': w,
                        'height': h,
                    })
                except Exception:
                    logger.debug('Failed to extract image from %s', filename)
                    continue
    except Exception:
        logger.debug('Failed to iterate document relationships for %s', filename)

    return {
        'filename': filename,
        'page_count': page_count,
        'text': full_text,
        'images': images,
        'metadata': metadata,
    }
